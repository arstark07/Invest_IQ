#!/usr/bin/env python3
"""
Convert Markdown report to Word DOCX using python-docx.
Requires: pip install python-docx
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_markdown(filepath):
    """Read Markdown file and return content."""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def create_word_doc(md_content, output_path):
    """Convert Markdown to Word document."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Full-Stack AI Finance Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Project Report — December 2025')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.italic = True
    
    # Split content into lines
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines at start
        if not line.strip():
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Handle lists
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        
        # Handle code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text, style='Normal')
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        
        # Handle tables
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Simple table parsing
            header = [cell.strip() for cell in line.split('|')[1:-1]]
            i += 2  # Skip separator
            
            rows = []
            while i < len(lines) and '|' in lines[i]:
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                rows.append(row)
                i += 1
            
            if header:
                table = doc.add_table(rows=1 + len(rows), cols=len(header))
                table.style = 'Light Grid Accent 1'
                
                # Add header
                hdr_cells = table.rows[0].cells
                for j, cell_text in enumerate(header):
                    hdr_cells[j].text = cell_text
                
                # Add rows
                for r_idx, row in enumerate(rows, start=1):
                    row_cells = table.rows[r_idx].cells
                    for c_idx, cell_text in enumerate(row):
                        row_cells[c_idx].text = cell_text
            continue
        
        # Handle horizontal rules
        elif line.strip() in ['---', '***', '___']:
            doc.add_paragraph()
        
        # Regular paragraphs
        elif line.strip():
            doc.add_paragraph(line)
        
        i += 1
    
    # Save document
    doc.save(output_path)
    print(f"✓ Document created: {output_path}")

if __name__ == '__main__':
    script_dir = Path(__file__).parent.parent
    md_file = script_dir / 'report' / 'report.md'
    out_file = script_dir / 'report' / 'AI_Finance_Project_Report.docx'
    
    if not md_file.exists():
        print(f"Error: {md_file} not found")
        sys.exit(1)
    
    print(f"Reading: {md_file}")
    content = read_markdown(md_file)
    
    print(f"Creating: {out_file}")
    create_word_doc(content, out_file)
    
    print("Done!")
